"""
临时脚本：解析文案模板文件
"""

from docx import Document
import re

def parse_template_file(file_path):
    """读取并解析Word模板文件"""
    doc = Document(file_path)

    print("="*60)
    print("文案模板文件内容预览")
    print("="*60)

    full_text = []
    for i, para in enumerate(doc.paragraphs[:100]):  # 只读前100段
        text = para.text.strip()
        if text:
            full_text.append(text)
            print(f"{i+1}. {text[:80]}{'...' if len(text) > 80 else ''}")

    print("\n" + "="*60)
    print(f"总段落数: {len(doc.paragraphs)}")
    print("="*60)

    # 尝试识别款式分类
    print("\n正在分析款式分类...")
    categories = []
    for text in full_text[:50]:
        # 查找可能的款式关键词
        keywords = ['羽绒', '马甲', '毛衣', '皮衣', '外套', '裤子', '羊毛', '针织']
        for kw in keywords:
            if kw in text:
                print(f"找到关键词 '{kw}': {text[:60]}...")
                break

    return full_text

if __name__ == '__main__':
    template_path = '/Users/haihui/dou/文案.docx'
    parse_template_file(template_path)
